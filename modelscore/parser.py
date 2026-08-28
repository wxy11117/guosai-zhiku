"""安全的只读材料解析器，不执行用户代码。"""

from __future__ import annotations

import re
import zipfile
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from pypdf import PdfReader


MAX_TEXT_CHARS = 120_000
MAX_ZIP_FILES = 500
MAX_UNCOMPRESSED = 80 * 1024 * 1024


@dataclass
class EvidencePage:
    page: int
    text: str


@dataclass
class ParsedMaterial:
    name: str
    kind: str
    pages: list[EvidencePage] = field(default_factory=list)
    inventory: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)

    @property
    def text(self) -> str:
        return "\n".join(page.text for page in self.pages)[:MAX_TEXT_CHARS]


def _clean(text: str) -> str:
    return re.sub(r"[ \t]+", " ", re.sub(r"\r\n?", "\n", text or "")).strip()


def _parse_pdf(path: Path, kind: str) -> ParsedMaterial:
    reader = PdfReader(str(path))
    pages = [EvidencePage(i + 1, _clean(page.extract_text() or "")) for i, page in enumerate(reader.pages)]
    char_count = sum(len(page.text) for page in pages)
    empty_pages = sum(not page.text for page in pages)
    result = ParsedMaterial(path.name, kind, pages, metadata={
        "format": "pdf", "page_count": len(pages), "page_locator": "exact",
        "char_count": char_count, "empty_pages": empty_pages,
        "text_coverage": round((len(pages) - empty_pages) / max(len(pages), 1), 3),
        "encrypted": bool(reader.is_encrypted),
    })
    if char_count < max(80, len(pages) * 20):
        result.warnings.append("PDF 文本覆盖率过低，可能是扫描件；本机未配置 OCR，当前评分置信度会降低。")
    return result


def _parse_docx(path: Path, kind: str) -> ParsedMaterial:
    doc = Document(str(path))
    blocks = [_clean(p.text) for p in doc.paragraphs if _clean(p.text)]
    pages = [EvidencePage(i // 45 + 1, "\n".join(blocks[i:i + 45])) for i in range(0, len(blocks), 45)]
    tables = len(doc.tables)
    return ParsedMaterial(path.name, kind, pages or [EvidencePage(1, "")], metadata={
        "format": "docx", "page_count": max(1, len(pages)), "page_locator": "estimated",
        "char_count": sum(len(block) for block in blocks), "paragraph_count": len(blocks), "table_count": tables,
    })


def _safe_zip_inventory(path: Path) -> tuple[list[str], list[str]]:
    warnings: list[str] = []
    inventory: list[str] = []
    with zipfile.ZipFile(path) as archive:
        infos = archive.infolist()
        if len(infos) > MAX_ZIP_FILES:
            warnings.append(f"压缩包文件过多，仅检查前 {MAX_ZIP_FILES} 项。")
        total = 0
        for info in infos[:MAX_ZIP_FILES]:
            name = info.filename.replace("\\", "/")
            if name.startswith("/") or "../" in name:
                warnings.append(f"跳过不安全路径：{name}")
                continue
            total += info.file_size
            if total > MAX_UNCOMPRESSED:
                warnings.append("压缩包解压后体积超过安全检查上限。")
                break
            inventory.append(name)
    return inventory, warnings


def parse_material(path: Path, kind: str) -> ParsedMaterial:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _parse_pdf(path, kind)
    if suffix == ".docx":
        return _parse_docx(path, kind)
    if suffix in {".txt", ".md", ".py", ".m", ".csv", ".json", ".ipynb"}:
        text = _clean(path.read_text(encoding="utf-8", errors="replace"))[:MAX_TEXT_CHARS]
        return ParsedMaterial(path.name, kind, [EvidencePage(1, text)], metadata={
            "format": suffix.lstrip("."), "page_count": 1, "page_locator": "logical", "char_count": len(text),
        })
    if suffix == ".zip":
        inventory, warnings = _safe_zip_inventory(path)
        return ParsedMaterial(path.name, kind, [EvidencePage(1, "\n".join(inventory))], inventory, warnings, {
            "format": "zip", "page_count": 1, "page_locator": "inventory", "file_count": len(inventory),
        })
    return ParsedMaterial(path.name, kind, warnings=[f"暂不支持解析 {suffix or '未知'} 格式，仅记录文件信息。"], metadata={"format": suffix.lstrip(".")})


def analyze_structure(material: ParsedMaterial) -> dict:
    text = material.text
    section_patterns = ["摘要", "关键词", "问题重述", "模型假设", "符号说明", "模型建立", "模型求解", "结果分析", "敏感性分析", "误差分析", "模型评价", "结论", "参考文献"]
    sections = [name for name in section_patterns if name in text]
    formulas = len(re.findall(r"(?:式\s*[（(]?\d+|[=≤≥]\s*[-+]?\d|\\begin\{equation\})", text, re.I))
    figures = len(re.findall(r"(?:图|figure)\s*[（(]?\d+", text, re.I))
    tables = len(re.findall(r"(?:表|table)\s*[（(]?\d+", text, re.I))
    metadata = dict(material.metadata)
    metadata.update({"sections": sections, "section_count": len(sections), "formula_mentions": formulas, "figure_mentions": figures, "table_mentions": tables})
    return metadata


def find_evidence(material: ParsedMaterial, keywords: tuple[str, ...]) -> dict | None:
    for page in material.pages:
        lower = page.text.lower()
        for keyword in keywords:
            pos = lower.find(keyword.lower())
            if pos >= 0:
                start = max(0, pos - 55)
                end = min(len(page.text), pos + len(keyword) + 85)
                quote = page.text[start:end].replace("\n", " ")
                return {"file": material.name, "page": page.page, "page_locator": material.metadata.get("page_locator", "unknown"), "quote": quote}
    return None
